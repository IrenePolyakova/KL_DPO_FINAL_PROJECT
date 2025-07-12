import streamlit as st
import os
import zipfile
import tempfile
from transformers import MarianMTModel, MarianTokenizer
from docx import Document
from io import BytesIO
import pandas as pd

st.set_page_config(page_title="Сравнение переводов MarianMT", layout="wide")
st.title("📄 Сравнение переводов с MarianMT (только локальные модели)")
st.markdown("Загрузите документ и ZIP-файлы моделей для сравнения.")

# Выбор моделей (локальные только)
num_models = st.number_input("Сколько моделей вы хотите загрузить?", min_value=1, max_value=5, value=1)
model_names = [f"Model_{i+1}" for i in range(num_models)]

# Загрузка zip-файлов моделей
uploaded_finetuned_paths = {}
st.markdown("### 🔄 Загрузите ZIP-файлы моделей")
for model in model_names:
    uploaded_zip = st.file_uploader(f"ZIP-файл модели {model}", type="zip", key=f"upload_{model}")
    if uploaded_zip is not None:
        with tempfile.TemporaryDirectory() as tmpdir:
            zip_path = os.path.join(tmpdir, "model.zip")
            with open(zip_path, "wb") as f:
                f.write(uploaded_zip.read())

            extract_path = os.path.join("uploaded_finetuned_models", model)
            os.makedirs(extract_path, exist_ok=True)

            with zipfile.ZipFile(zip_path, "r") as zip_ref:
                zip_ref.extractall(extract_path)

            uploaded_finetuned_paths[model] = extract_path
        st.success(f"Модель {model} успешно загружена и распакована.")

# Загрузка DOCX
uploaded_file = st.file_uploader("Загрузите DOCX-документ", type=["docx"])

if uploaded_file and uploaded_finetuned_paths:
    doc = Document(uploaded_file)
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    if not paragraphs:
        st.warning("Документ не содержит текста для перевода.")
    else:
        st.info(f"Найдено абзацев для перевода: {len(paragraphs)}")

        translations = {model: [] for model in uploaded_finetuned_paths}
        translations["Original"] = paragraphs

        for model_name, model_path in uploaded_finetuned_paths.items():
            st.write(f"🔄 Перевод с помощью {model_name}...")

            try:
                tokenizer = MarianTokenizer.from_pretrained(model_path)
                model = MarianMTModel.from_pretrained(model_path)
            except Exception as e:
                st.error(f"Ошибка загрузки модели {model_name}: {e}")
                continue

            translated = []
            for p in paragraphs:
                inputs = tokenizer(p, return_tensors="pt", padding=True, truncation=True)
                outputs = model.generate(**inputs)
                translated_text = tokenizer.decode(outputs[0], skip_special_tokens=True)
                translated.append(translated_text)
            translations[model_name] = translated

        # Таблица результатов
        st.markdown("## 📊 Сравнение переводов")
        df = pd.DataFrame(translations)
        st.dataframe(df, use_container_width=True)

        # Выгрузка результатов
        st.markdown("### 📥 Скачать результаты")
        export_format = st.selectbox("Формат выгрузки", ["DOCX", "CSV"])
        if export_format == "CSV":
            csv = df.to_csv(index=False).encode("utf-8")
            st.download_button("Скачать CSV", data=csv, file_name="translations.csv", mime="text/csv")
        else:
            out_doc = Document()
            table = out_doc.add_table(rows=1, cols=len(df.columns))
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(df.columns):
                hdr_cells[i].text = col
            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = val
            docx_io = BytesIO()
            out_doc.save(docx_io)
            st.download_button("Скачать DOCX", data=docx_io.getvalue(), file_name="translations.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
