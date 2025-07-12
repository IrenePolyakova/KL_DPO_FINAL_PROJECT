import io
import re
import logging
from typing import List, Dict, Tuple, Optional, Iterator, BinaryIO
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table, _Cell
from docx.section import Section
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl

logger = logging.getLogger(__name__)

def extract_content_from_docx(docx_file: BinaryIO) -> Dict:
    """
    Извлекает содержимое из DOCX файла
    
    Args:
        docx_file: Файловый объект DOCX
        
    Returns:
        Dict: Словарь с извлеченным содержимым:
            - sentences: список предложений
            - tables: список таблиц
            - metadata: метаданные документа
    """
    try:
        doc = Document(docx_file)
        sentences = []
        tables = []
        metadata = {
            'title': doc.core_properties.title or '',
            'author': doc.core_properties.author or '',
            'created': doc.core_properties.created,
            'modified': doc.core_properties.modified,
            'language': doc.core_properties.language or ''
        }
        
        # Извлекаем предложения из основного текста
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                sentences.extend(split_into_sentences(paragraph.text))
        
        # Извлекаем таблицы
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                    row_data.append(cell_text)
                if any(cell.strip() for cell in row_data):
                    table_data.append(row_data)
            if table_data:
                tables.append(table_data)
        
        # Извлекаем текст из колонтитулов
        for section in doc.sections:
            for header in [section.header, section.first_page_header]:
                if header:
                    for paragraph in header.paragraphs:
                        if paragraph.text.strip():
                            sentences.extend(split_into_sentences(paragraph.text))
            
            for footer in [section.footer, section.first_page_footer]:
                if footer:
                    for paragraph in footer.paragraphs:
                        if paragraph.text.strip():
                            sentences.extend(split_into_sentences(paragraph.text))
        
        return {
            'sentences': sentences,
            'tables': tables,
            'metadata': metadata
        }
        
    except Exception as e:
        logger.error(f"Ошибка при извлечении содержимого DOCX: {str(e)}")
        return {'sentences': [], 'tables': [], 'metadata': {}}

def split_into_sentences(text: str) -> List[str]:
    """
    Разбивает текст на предложения с учетом различных случаев
    
    Args:
        text: Исходный текст
        
    Returns:
        List[str]: Список предложений
    """
    if not text or not text.strip():
        return []
    
    try:
        # Сохраняем пробелы в начале и конце
        leading_space = ' ' * (len(text) - len(text.lstrip()))
        trailing_space = ' ' * (len(text) - len(text.rstrip()))
        
        # Регулярное выражение для разбиения на предложения
        # Учитываем:
        # - Стандартные знаки конца предложения (. ! ?)
        # - Многоточие (...)
        # - Сокращения (Mr., Dr., т.д., т.п.)
        # - Инициалы (A.B., Ю.А.)
        # - Числа с точкой (1.5, 2.0)
        pattern = r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<![A-Z]\.)(?<!\d\.)(?<=\.|\?|!|\.\.\.|。|！|？|।|؟|।|。)\s+'
        
        sentences = re.split(pattern, text.strip())
        
        # Восстанавливаем пробелы
        if sentences:
            sentences[0] = leading_space + sentences[0]
            sentences[-1] = sentences[-1] + trailing_space
        
        return [s for s in sentences if s.strip()]
        
    except Exception as e:
        logger.warning(f"Ошибка при разбиении на предложения: {str(e)}")
        return [text]

def create_translated_docx(
    src_docx: BinaryIO,
    translations: Dict[str, List[str]],
    output_path: Optional[str] = None
) -> bytes:
    """
    Создает DOCX с переводами, сохраняя форматирование
    
    Args:
        src_docx: Исходный DOCX файл
        translations: Словарь с переводами для разных систем
        output_path: Путь для сохранения файла (опционально)
        
    Returns:
        bytes: Содержимое DOCX файла
    """
    try:
        src_docx.seek(0)
        doc = Document(src_docx)
        
        # Создаем отдельные секции для каждой системы перевода
        for system_name, trans_sentences in translations.items():
            # Добавляем заголовок системы
            heading = doc.add_heading(f'Перевод: {system_name}', level=1)
            heading.style.font.size = Pt(14)
            heading.style.font.color.rgb = RGBColor(0, 0, 0)
            
            # Копируем структуру и содержимое документа
            trans_iter = iter(trans_sentences)
            
            # Обрабатываем основной текст
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    process_paragraph(paragraph, trans_iter)
            
            # Обрабатываем таблицы
            for table in doc.tables:
                process_table(table, trans_iter)
            
            # Обрабатываем колонтитулы
            for section in doc.sections:
                process_section(section, trans_iter)
        
        # Сохраняем результат
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        if output_path:
            with open(output_path, 'wb') as f:
                f.write(output.getvalue())
        
        return output.getvalue()
        
    except Exception as e:
        logger.error(f"Ошибка при создании переведенного DOCX: {str(e)}")
        raise

def process_paragraph(paragraph: Paragraph, trans_iter: Iterator[str]) -> None:
    """Обрабатывает параграф, сохраняя форматирование"""
    try:
        # Собираем все runs с текстом
        text_runs = [run for run in paragraph.runs if run.text.strip()]
        if not text_runs:
            return
        
        # Получаем следующий перевод
        try:
            trans_text = next(trans_iter)
        except StopIteration:
            logger.warning("Закончились переводы до окончания документа")
            return
        
        # Сохраняем стили всех runs
        run_styles = []
        for run in text_runs:
            style = {
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'strike': run.font.strike,
                'subscript': run.font.subscript,
                'superscript': run.font.superscript,
                'size': run.font.size,
                'name': run.font.name,
                'color': run.font.color.rgb if run.font.color else None,
                'highlight_color': run.font.highlight_color,
                'language': run.font.language
            }
            run_styles.append(style)
        
        # Очищаем параграф
        for run in paragraph.runs:
            run.text = ""
        
        # Создаем новый run с переводом
        new_run = paragraph.add_run(trans_text)
        
        # Применяем стили первого run
        if run_styles:
            style = run_styles[0]
            new_run.bold = style['bold']
            new_run.italic = style['italic']
            new_run.underline = style['underline']
            new_run.font.strike = style['strike']
            new_run.font.subscript = style['subscript']
            new_run.font.superscript = style['superscript']
            if style['size']:
                new_run.font.size = style['size']
            if style['name']:
                new_run.font.name = style['name']
            if style['color']:
                new_run.font.color.rgb = style['color']
            if style['highlight_color']:
                new_run.font.highlight_color = style['highlight_color']
            if style['language']:
                new_run.font.language = style['language']
                
    except Exception as e:
        logger.error(f"Ошибка при обработке параграфа: {str(e)}")

def process_table(table: Table, trans_iter: Iterator[str]) -> None:
    """Обрабатывает таблицу"""
    try:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        process_paragraph(paragraph, trans_iter)
    except Exception as e:
        logger.error(f"Ошибка при обработке таблицы: {str(e)}")

def process_section(section: Section, trans_iter: Iterator[str]) -> None:
    """Обрабатывает секцию документа"""
    try:
        # Обрабатываем колонтитулы
        for header in [section.header, section.first_page_header]:
            if header:
                for paragraph in header.paragraphs:
                    if paragraph.text.strip():
                        process_paragraph(paragraph, trans_iter)
        
        for footer in [section.footer, section.first_page_footer]:
            if footer:
                for paragraph in footer.paragraphs:
                    if paragraph.text.strip():
                        process_paragraph(paragraph, trans_iter)
                        
    except Exception as e:
        logger.error(f"Ошибка при обработке секции: {str(e)}")

def align_sentences(src_sentences: List[str], ref_sentences: List[str]) -> Tuple[List[str], List[str]]:
    """
    Выравнивает предложения между исходным и эталонным текстом
    
    Args:
        src_sentences: Исходные предложения
        ref_sentences: Эталонные предложения
        
    Returns:
        Tuple[List[str], List[str]]: Выровненные списки предложений
    """
    if not src_sentences or not ref_sentences:
        return [], []
    
    try:
        # Нормализуем предложения
        src_norm = [s.strip().lower() for s in src_sentences]
        ref_norm = [r.strip().lower() for r in ref_sentences]
        
        # Создаем матрицу схожести
        import numpy as np
        from difflib import SequenceMatcher
        
        similarity_matrix = np.zeros((len(src_norm), len(ref_norm)))
        for i, src in enumerate(src_norm):
            for j, ref in enumerate(ref_norm):
                similarity_matrix[i, j] = SequenceMatcher(None, src, ref).ratio()
        
        # Находим наилучшие соответствия
        aligned_src = []
        aligned_ref = []
        used_ref = set()
        
        for i, src in enumerate(src_sentences):
            best_match = -1
            best_score = 0.5  # Порог схожести
            
            for j, ref in enumerate(ref_sentences):
                if j not in used_ref and similarity_matrix[i, j] > best_score:
                    best_match = j
                    best_score = similarity_matrix[i, j]
            
            if best_match != -1:
                aligned_src.append(src)
                aligned_ref.append(ref_sentences[best_match])
                used_ref.add(best_match)
            else:
                aligned_src.append(src)
                aligned_ref.append("")
        
        # Добавляем оставшиеся эталонные предложения
        for j, ref in enumerate(ref_sentences):
            if j not in used_ref:
                aligned_src.append("")
                aligned_ref.append(ref)
        
        return aligned_src, aligned_ref
        
    except Exception as e:
        logger.error(f"Ошибка при выравнивании предложений: {str(e)}")
        return src_sentences, ref_sentences