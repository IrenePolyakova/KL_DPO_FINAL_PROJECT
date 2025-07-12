import io
import re
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def extract_sentences_from_docx(docx_file):
    """Извлекает предложения из DOCX файла"""
    doc = Document(docx_file)
    sentences = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            # Разбиваем на предложения с сохранением пробелов
            sents = split_into_sentences(paragraph.text)
            sentences.extend(sents)
    
    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        sents = split_into_sentences(paragraph.text)
                        sentences.extend(sents)
    
    # Обработка колонтитулов
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header is not None:
                for paragraph in header.paragraphs:
                    if paragraph.text.strip():
                        sents = split_into_sentences(paragraph.text)
                        sentences.extend(sents)
        
        for footer in [section.footer, section.first_page_footer]:
            if footer is not None:
                for paragraph in footer.paragraphs:
                    if paragraph.text.strip():
                        sents = split_into_sentences(paragraph.text)
                        sentences.extend(sents)
    
    return sentences

def split_into_sentences(text):
    """Разбивает текст на предложения, сохраняя пробелы и знаки препинания"""
    # Улучшенное разбиение: учитываем многоточие, восклицательные и вопросительные знаки
    sentence_endings = re.compile(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<![A-Z]\.)(?<=\.|\?|!|…)\s+')
    sentences = sentence_endings.split(text)
    # Убираем пустые строки
    return [sent for sent in sentences if sent.strip()]

def preserve_formatting_translation(src_docx, translated_sentences):
    """Создает DOCX с переведенным текстом, сохраняя форматирование, изображения и структуру"""
    src_docx.seek(0)
    doc = Document(src_docx)
    trans_iter = iter(translated_sentences)
    
    # Обработка параграфов в основном тексте
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            process_paragraph(paragraph, trans_iter)
    
    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        process_paragraph(paragraph, trans_iter)
    
    # Обработка колонтитулов
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header is not None:
                for paragraph in header.paragraphs:
                    if paragraph.text.strip():
                        process_paragraph(paragraph, trans_iter)
        
        for footer in [section.footer, section.first_page_footer]:
            if footer is not None:
                for paragraph in footer.paragraphs:
                    if paragraph.text.strip():
                        process_paragraph(paragraph, trans_iter)
    
    # Сохраняем документ
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()

def process_paragraph(paragraph, trans_iter):
    """Обрабатывает параграф, заменяя текст с сохранением форматирования"""
    # Собираем все runs, которые содержат текст
    runs = [run for run in paragraph.runs if run.text.strip()]
    if not runs:
        return
        
    # Заменяем текст в каждом run на перевод
    for run in runs:
        try:
            trans_text = next(trans_iter)
        except StopIteration:
            break
            
        # Сохраняем стиль run
        original_font = run.font
        run_props = {
            'bold': original_font.bold,
            'italic': original_font.italic,
            'underline': original_font.underline,
            'color': original_font.color.rgb if original_font.color and original_font.color.rgb else None,
            'size': original_font.size,
            'name': original_font.name,
        }
        
        # Заменяем текст
        run.text = trans_text
        
        # Восстанавливаем стиль
        run.font.bold = run_props['bold']
        run.font.italic = run_props['italic']
        run.font.underline = run_props['underline']
        if run_props['color']:
            run.font.color.rgb = run_props['color']
        if run_props['size']:
            run.font.size = run_props['size']
        if run_props['name']:
            run.font.name = run_props['name']
    
    # Удаляем пустые runs
    for run in paragraph.runs:
        if not run.text.strip():
            run.text = ""